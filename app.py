# app.py  – auto‑downloads spaCy model if missing
import io
import fitz  # PyMuPDF
import streamlit as st
import pdfplumber
from docx import Document
import spacy

# ▶︎ Ensure the spaCy model is available even on fresh Streamlit Cloud builds
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    from spacy.cli import download
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

EDU = [
    "university", "college", "institute", "school", "faculty", "academy"
]

def is_edu(text):
    return any(k in text.lower() for k in EDU)

def detect_entities(text: str):
    """Return list of PERSON names + education entities to redact."""
    targets = set()
    for ent in nlp(text).ents:
        if ent.label_ == "PERSON" or is_edu(ent.text):
            targets.add(ent.text.strip())
    return list(targets)

# ---------- DOCX helpers ---------- #

def extract_docx_text(bytes_data):
    doc = Document(io.BytesIO(bytes_data))
    return "\n".join(p.text for p in doc.paragraphs)

def redact_docx(bytes_data, targets):
    src = Document(io.BytesIO(bytes_data))
    dst = Document()
    for para in src.paragraphs:
        line = para.text
        for t in targets:
            line = line.replace(t, "█" * len(t))
        dst.add_paragraph(line)
    out = io.BytesIO(); dst.save(out)
    return out.getvalue()

# ---------- PDF helpers ---------- #

def redact_pdf(bytes_data, targets):
    pdf = fitz.open(stream=bytes_data, filetype="pdf")
    for page in pdf:
        for t in targets:
            for rect in page.search_for(t):
                page.add_redact_annot(rect, fill=(0, 0, 0))
        page.apply_redactions()
    out = io.BytesIO(); pdf.save(out)
    return out.getvalue()

# ---------- Streamlit UI ---------- #
st.set_page_config(page_title="CV Anonymiser", layout="centered")
st.title("🕵️ CV Anonymiser – Zero‑Bias Redaction")

st.markdown(
    "Upload a **PDF** or **DOCX** CV. The app finds personal names and education institutions and returns a file with blacked‑out text that is no longer selectable."
)

file = st.file_uploader("Upload CV", type=["pdf", "docx"])

if file:
    ext = file.name.split(".")[-1].lower()
    data = file.read()

    # Extract plain text for entity detection
    if ext == "pdf":
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    else:
        text = extract_docx_text(data)

    targets = detect_entities(text)
    st.success(f"Detected {len(targets)} items to redact.")
    st.write(targets)

    if st.button("Redact & Download"):
        redacted = redact_pdf(data, targets) if ext == "pdf" else redact_docx(data, targets)
        outfile = f"redacted_cv.{ext if ext=='pdf' else 'docx'}"
        st.download_button("Download Redacted CV", redacted, file_name=outfile, mime="application/octet-stream")
